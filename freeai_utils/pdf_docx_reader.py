import os
import fitz # pip install pymupdf
from pypdf import PdfReader # need pip install pypdf
import pdfplumber # need pip install pdfplumber
from typing import Optional
import logging
logging.getLogger("pdfminer").setLevel(logging.ERROR) #stop the pdfminer from displaying logs that just info or debug
from freeai_utils.log_set_up import setup_logging
try:
    from docx import Document # need pip install python-docx
except ImportError:
    Document = None

class PDF_DOCX_Reader:
    def __init__(self, start_page: int = 0, last_page: Optional[int] = None) -> None:
        # Initialize with optional page range.
        # Supports PDF and DOCX formats.
        # param first_page: zero-based index of first page to process
        # param last_page: zero-based index of last page, None means all pages
        self.__enforce_type(start_page, int, "first_page")
        self.__enforce_type(last_page, (int, type(None)), "last_page")
        
        self.first_page = start_page
        self.last_page = last_page
        #for logging only this class rather
        self.logger = setup_logging(self.__class__.__name__)
        self.logger.info(f"Initialize successfully")

    def extract_all_text(self, file_path: str = None, first_page: Optional[int] = None, last_page: Optional[int] = None) -> str:
        # Extract all text into a single string.
        # Tries pypdf for PDF, Document for DOCX, and falls back to fitz for PDFs.
        
        #check type
        self.__enforce_type(first_page, (int, type(None)), "first_page")
        self.__enforce_type(last_page, (int, type(None)), "last_page")
        
        self.__enforce_type(file_path, str, file_path)
        ext = self.__isSupported(file_path) #extract the end (file type)
        
        #after config, you can still choose the start and end page for specific one
        fp = first_page if first_page is not None else self.first_page
        lp = last_page if last_page is not None else self.last_page

        self.logger.info(f"Detect file type: {ext}")
        
        # DOCX
        if ext == '.docx':
            if Document is None:
                raise ImportError("python-docx is required to read DOCX files.")
            doc = Document(file_path)
            text = []
            try:
                for para in doc.paragraphs:
                    if para.text:
                        text.append(para.text)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text.append(cell.text)
                return '\n'.join(text)
            except Exception as e:
                self.logger.error(f"Fail to get text from file {file_path} with error {e}")
                return ""

        # PDF
        text = []
        try:
            reader = PdfReader(file_path)
            pages = reader.pages
            total = len(pages)
            end = lp + 1 if lp is not None else total
            for i in range(fp, min(end, total)):
                page = pages[i]
                txt = page.extract_text() or ''
                text.append(txt)
            return '\n'.join(text)
           
        except Exception:
            # fallback to fitz if can't use pypdf
            self.logger.error(f"Fail to use pypdf, fallback to fitz to extract file")
            doc = fitz.open(file_path)
            total = doc.page_count
            end = lp + 1 if lp is not None else total
            try:
                for i in range(fp, min(end, total)):
                    page = doc.load_page(i)
                    text.append(page.get_text())
                return '\n'.join(text)
            except Exception as e:
                self.logger.error(f"Fail to get text from file {file_path} with error {e}")
                return ""

    def extract_ordered_text( self, file_path: str = None, first_page: Optional[int] = None, last_page: Optional[int] = None) -> str:
        # Extract text maintaining layout using pdfplumber for PDFs,
        # or fallback to fitz. DOCX behaves same as extract_all_text.
        
        #check type
        self.__enforce_type(first_page, (int, type(None)), "first_page")
        self.__enforce_type(last_page, (int, type(None)), "last_page")

        #path, file checking
        self.__enforce_type(file_path, str, "file_path")
        ext = self.__isSupported(file_path)
        
        fp = first_page if first_page is not None else self.first_page
        lp = last_page if last_page is not None else self.last_page
        
        # DOCX same as extract_all_text
        if ext == '.docx':
            #adjust here the code please
            try:
                from docx.oxml.text.paragraph import CT_P
                from docx.oxml.table import CT_Tbl
                from docx.text.paragraph import Paragraph
                from docx.table import Table
                doc = Document(file_path)
                parts = []
                for child in doc.element.body.iterchildren():
                    if isinstance(child, CT_P):
                        para = Paragraph(child, doc)
                        parts.append(para.text)
                    elif isinstance(child, CT_Tbl):
                        tbl = Table(child, doc)
                        # turn each row into a tab‑joined line
                        for row in tbl.rows:
                            parts.append("\t".join(cell.text for cell in row.cells))
                return "\n".join(parts)
            except Exception:
                self.logger.error("Fail to do extract_ordered_text for docx file, trying extract_all_text instead")
                return self.extract_all_text(file_path, fp, lp)
        
        # PDF using pdfplumber
        text = []
        try:
            with pdfplumber.open(file_path) as pdf:
                total = len(pdf.pages)
                end = lp + 1 if lp is not None else total
                for i in range(fp, min(end, total)):
                    page = pdf.pages[i]
                    # preserve layout spacing
                    txt = page.extract_text(x_tolerance=2, y_tolerance=2) or ''
                    text.append(txt)
            return '\n'.join(text)
        except Exception:
            # fallback to fitz
            self.logger.info(f"Fail to use pdflumber, fallback to fitz to extract file")
            return self.extract_all_text(file_path, fp, lp)

    def extract_images(self, file_path: str = None, folder_extract: str = "extracted_images", first_page: Optional[int] = None, last_page: Optional[int] = None) -> int:
        # Extract images from PDF using fitz (PyMuPDF).
        # Saves images to folder_extract, and returns count.
   
        #check type
        self.__enforce_type(first_page, (int, type(None)), "first_page")
        self.__enforce_type(last_page, (int, type(None)), "last_page")

        #path, file checking
        self.__enforce_type(file_path, str, "file_path")
        ext = self.__isSupported(file_path)
        
        fp = first_page if first_page is not None else self.first_page
        lp = last_page if last_page is not None else self.last_page

        os.makedirs(folder_extract, exist_ok=True) #folder create or use existence one
        count = 0

        self.logger.info(f"Detect file type: {ext}")
        
        if ext == '.pdf':
            doc = fitz.open(file_path)
            total = doc.page_count
            end = lp + 1 if lp is not None else total
            for i in range(fp, min(end, total)):
                page = doc.load_page(i)
                images = page.get_images(full=True)
                for img_index, img in enumerate(images):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    img_bytes = base_image["image"]
                    extn = base_image.get("ext", "png")
                    img_name = f"page{i+1}_img{img_index+1}.{extn}"
                    out_path = os.path.join(folder_extract, img_name)
                    with open(out_path, 'wb') as f:
                        f.write(img_bytes)
                    count += 1
        else:
            # DOCX image extraction using related_parts
            if Document is None:
                raise ImportError("python-docx is required to extract images from DOCX files.")
            doc = Document(file_path)
            # collect all image parts
            image_parts = [part for part in doc.part.related_parts.values()
                           if hasattr(part, 'content_type') and part.content_type.startswith('image/')]
            for part in image_parts:
                img_bytes = part.blob
                img_name = os.path.basename(part.partname)
                out_path = os.path.join(folder_extract, img_name)
                with open(out_path, 'wb') as f:
                    f.write(img_bytes)
                count += 1

        self.logger.info(f"Complete extracted images in file {file_path} to folder {folder_extract}")
        return count #return number of images found in the file

    def __isSupported(self, file_path : str) -> str:
        if file_path:
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in ['.pdf', '.docx']:
                raise ValueError(f"Unsupported file type: {ext}")
            return ext
    
    def __enforce_type(self, value, expected_types, arg_name) -> None:
        if not isinstance(value, expected_types):
            expected_names = [t.__name__ for t in expected_types] if isinstance(expected_types, tuple) else [expected_types.__name__]
            expected_str = ", ".join(expected_names)
            raise TypeError(f"Argument '{arg_name}' must be of type {expected_str}, but received {type(value).__name__}")